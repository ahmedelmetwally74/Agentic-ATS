import os
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from llm_service import (
    rewrite_summary_for_cv,
    rewrite_experience_for_cv,
    rewrite_education_for_cv,
    rewrite_skills_for_cv,
    rewrite_projects_for_cv,
    rewrite_simple_list_for_cv,
)


REQUIRED_FILES = [
    "name.txt",
    "email.txt",
    "country code.txt",
    "mobile.txt",
    "title.txt",
    "education.txt",
    "skills.txt",
    "projects.txt",
]

OPTIONAL_FILES = [
    "linkedin link.txt",
    "github link.txt",
    "military status.txt",
    "summary.txt",
    "experience.txt",
    "courses.txt",
    "awards.txt",
    "activities.txt",
]


def _read_text_file(path: str) -> str:
    if not os.path.isfile(path):
        return ""
    with open(path, "r", encoding="utf-8") as f:
        return f.read().strip()


def has_content(text: str) -> bool:
    return bool(text and text.strip())


def read_cv_sections(sections_dir: str) -> Dict[str, str]:
    data = {}
    all_files = REQUIRED_FILES + OPTIONAL_FILES
    for file_name in all_files:
        data[file_name] = _read_text_file(os.path.join(sections_dir, file_name))
    return data


def validate_required_sections(data: Dict[str, str]) -> None:
    missing = [name for name in REQUIRED_FILES if not has_content(data.get(name, ""))]
    if missing:
        raise ValueError("Missing required files: " + ", ".join(missing))


def decide_section_order(data: Dict[str, str]) -> List[str]:
    order = ["header"]

    if has_content(data.get("summary.txt", "")):
        order.append("summary")

    if has_content(data.get("experience.txt", "")):
        order.extend(["experience", "education", "projects", "skills"])
    else:
        order.extend(["education", "projects", "skills"])

    if has_content(data.get("courses.txt", "")):
        order.append("courses")
    if has_content(data.get("awards.txt", "")):
        order.append("awards")
    if has_content(data.get("activities.txt", "")):
        order.append("activities")

    return order


def add_hyperlink(paragraph, text: str, url: str):
    """
    Add a clickable hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_margins(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def _add_section_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def _add_summary_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.strip())
    run.font.size = Pt(10)

def _add_experience_block(doc: Document, text: str):
    """
    Expected format:

    AI / ML Engineer Intern at Siemens (2025 - Present)
    - Bullet 1
    - Bullet 2

    Machine Learning Intern at Giza Systems (2024)
    - Bullet 1
    - Bullet 2
    """
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]

    for block in blocks:
        lines = [x.strip() for x in block.splitlines() if x.strip()]
        if not lines:
            continue

        # First line = role/company/date
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_after = Pt(0)
        role_run = role_p.add_run(lines[0])
        role_run.bold = True
        role_run.font.size = Pt(10.5)

        # Remaining lines = bullets
        for bullet_line in lines[1:]:
            clean = bullet_line.lstrip("-• ").strip()
            if clean:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)
                run = bp.add_run(clean)
                run.font.size = Pt(10)

def _add_education_block(doc: Document, text: str):
    """
    Expected format:

    Bachelor's Degree in Information Technology
    Faculty of Computers and Information
    Mansoura University
    Graduation Project Grade: Excellent
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return

    # First line = degree
    deg_p = doc.add_paragraph()
    deg_p.paragraph_format.space_after = Pt(0)
    deg_run = deg_p.add_run(lines[0])
    deg_run.bold = True
    deg_run.font.size = Pt(10.5)

    # Remaining lines = details
    for line in lines[1:]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10)

def _add_education_block(doc: Document, text: str):
    """
    Expected format:

    Bachelor's Degree in Information Technology
    Faculty of Computers and Information, Mansoura University
    Graduation Project Grade: Excellent
    """
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    if not lines:
        return

    line1 = doc.add_paragraph()
    line1.paragraph_format.space_after = Pt(0)
    run1 = line1.add_run(lines[0])
    run1.bold = True
    run1.font.size = Pt(10.5)

    for line in lines[1:]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10)

def _add_skills_block(doc: Document, text: str):
    """
    Expected format:

    Programming Languages: Python, C, C++
    Machine Learning & Deep Learning: scikit-learn, TensorFlow, PyTorch
    Data Visualization: Pandas, Plotly
    """
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10)

def _add_bullets_from_text(doc: Document, text: str):
    """
    Turn lines into bullets.
    If lines already start with '-', remove it and keep bullet style.
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        clean = line.lstrip("-• ").strip()
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(clean)
        run.font.size = Pt(10)

def _clean_projects_output(text: str) -> str:
    if not text:
        return ""

    bad_markers = [
        "\nRules:",
        "\nOutput format must be exactly like this:",
        "\nAnother Project Title",
        "\nProject Title",
    ]

    cleaned = text.strip()
    for marker in bad_markers:
        pos = cleaned.find(marker)
        if pos != -1:
            cleaned = cleaned[:pos].strip()

    return cleaned

def _add_projects_block(doc: Document, text: str):
    """
    Expected text format after rewrite:
    Project Title
    - bullet
    - bullet

    Another Project
    - bullet
    """
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]

    for block in blocks:
        lines = [x.strip() for x in block.splitlines() if x.strip()]
        if not lines:
            continue

        title = lines[0]
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(0)
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(10.5)

        for bullet_line in lines[1:]:
            clean = bullet_line.lstrip("-• ").strip()
            if clean:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)
                run = bp.add_run(clean)
                run.font.size = Pt(10)


def generate_cv_docx(data: Dict[str, str], output_path: str) -> str:
    doc = Document()
    _set_margins(doc)

    # ---------------- Header ----------------
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data["name.txt"])
    name_run.bold = True
    name_run.font.size = Pt(18)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(data["title.txt"])
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_p.paragraph_format.space_after = Pt(4)

    line1 = doc.add_paragraph()
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1.paragraph_format.space_after = Pt(0)

    email_run = line1.add_run(data["email.txt"])
    email_run.font.size = Pt(10.5)

    military_status = data.get("military status.txt", "").strip()
    if military_status:
        sep = line1.add_run(" | ")
        sep.font.size = Pt(10.5)
        ms = line1.add_run(f"Military Service: {military_status}")
        ms.font.size = Pt(10.5)

    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line2.paragraph_format.space_after = Pt(6)

    linkedin = data.get("linkedin link.txt", "").strip()
    github = data.get("github link.txt", "").strip()
    phone = f"({data['country code.txt'].strip()}){data['mobile.txt'].strip()}"

    added_any = False
    if linkedin:
        add_hyperlink(line2, "LinkedIn", linkedin)
        added_any = True

    if github:
        if added_any:
            line2.add_run(" | ")
        add_hyperlink(line2, "GitHub", github)
        added_any = True

    if added_any:
        line2.add_run(" | ")

    phone_run = line2.add_run(phone)
    phone_run.font.size = Pt(10.5)

    # ---------------- Dynamic sections ----------------
    order = decide_section_order(data)

    for section_name in order:
        if section_name == "header":
            continue

        if section_name == "summary":
            raw = data.get("summary.txt", "")
            if has_content(raw):
                _add_section_title(doc, "Summary")
                polished = rewrite_summary_for_cv(raw)
                _add_summary_block(doc, polished)

        elif section_name == "experience":
            raw = data.get("experience.txt", "")
            if has_content(raw):
                _add_section_title(doc, "Experience")
                polished = rewrite_experience_for_cv(raw)
                _add_experience_block(doc, polished)

        elif section_name == "education":
            raw = data.get("education.txt", "")
            _add_section_title(doc, "Education")
            polished = rewrite_education_for_cv(raw)
            _add_education_block(doc, polished)

        elif section_name == "projects":
            raw = data.get("projects.txt", "")
            _add_section_title(doc, "Projects")
            polished = rewrite_projects_for_cv(raw)
            polished = _clean_projects_output(polished)
            _add_projects_block(doc, polished)

        elif section_name == "skills":
            raw = data.get("skills.txt", "")
            _add_section_title(doc, "Skills")
            polished = rewrite_skills_for_cv(raw)
            _add_skills_block(doc, polished)

        elif section_name == "courses":
            raw = data.get("courses.txt", "")
            if has_content(raw):
                _add_section_title(doc, "Courses & Certificates")
                polished = rewrite_simple_list_for_cv("courses", raw)
                _add_bullets_from_text(doc, polished)

        elif section_name == "awards":
            raw = data.get("awards.txt", "")
            if has_content(raw):
                _add_section_title(doc, "Awards")
                polished = rewrite_simple_list_for_cv("awards", raw)
                _add_bullets_from_text(doc, polished)

        elif section_name == "activities":
            raw = data.get("activities.txt", "")
            if has_content(raw):
                _add_section_title(doc, "Activities")
                polished = rewrite_simple_list_for_cv("activities", raw)
                _add_bullets_from_text(doc, polished)

    doc.save(output_path)
    return output_path


def generate_cv_from_sections(sections_dir: str, output_path: str = "generated_cv.docx") -> str:
    if not os.path.isdir(sections_dir):
        raise FileNotFoundError(f"Sections directory not found: {sections_dir}")

    data = read_cv_sections(sections_dir)
    validate_required_sections(data)

    final_path = generate_cv_docx(data, output_path)
    print(f"[SUCCESS] Generated CV saved to: {final_path}")
    return final_path