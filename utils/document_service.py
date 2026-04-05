"""
AgenticATS - Document Service
Core PDF and Word document extraction, generation, and redlining functions.
No Azure Document Intelligence / OCR dependencies.
"""

import fitz  # PyMuPDF
import logging
import os
import re
import uuid
import zipfile
import shutil
from pathlib import Path
from tempfile import mkdtemp

import PyPDF2
from docx import Document
from lxml import etree
from python_redlines.engines import XmlPowerToolsEngine

from utils.document_utils import normalize, flatten


# ---------------------------------------------------------------------------
# PDF Text Extraction
# ---------------------------------------------------------------------------

async def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).
    Sorts blocks by vertical (Y) coordinate to keep dates and left-aligned text on the same line.
    """
    return extract_text_from_pdf_sync(file_path)


def _extract_page_text_by_y(page) -> str:
    """
    Extracts text from a PyMuPDF page by grouping text spans into horizontal lines.
    This reconstructs the visual layout, ensuring right-aligned dates 
    stay on the same line as their left-aligned headings.
    """
    text_dict = page.get_text("dict")
    spans = []

    # Extract all text spans with their bounding boxes
    for block in text_dict.get("blocks", []):
        if block.get("type") == 0:  # Text block
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        # bbox: (x0, y0, x1, y1)
                        spans.append((span["bbox"], text))

    if not spans:
        return ""

    # Sort spans primarily by top Y-coordinate (y0)
    spans.sort(key=lambda x: x[0][1])

    lines = []
    current_line_spans = []
    
    # We define a "line" by tracking the bounding box of the current line
    # If a new span overlaps vertically with the current line, it belongs to it.
    current_y0 = spans[0][0][1]
    current_y1 = spans[0][0][3]

    for bbox, text in spans:
        span_y0, span_y1 = bbox[1], bbox[3]
        
        # Calculate vertical overlap
        overlap = max(0, min(current_y1, span_y1) - max(current_y0, span_y0))
        span_height = span_y1 - span_y0
        
        # If overlap is significant (e.g., > 50% of the span's height), it's the same line
        if span_height > 0 and (overlap / span_height) > 0.5:
            current_line_spans.append((bbox, text))
            # Expand the line's bounding box to encompass this span
            current_y0 = min(current_y0, span_y0)
            current_y1 = max(current_y1, span_y1)
        else:
            # Sort spans in the line by X-coordinate (left to right)
            current_line_spans.sort(key=lambda x: x[0][0])
            joined_text = "  ".join([s[1] for s in current_line_spans]).strip()
            if joined_text:
                lines.append(joined_text)
            
            # Start a new line
            current_line_spans = [(bbox, text)]
            current_y0 = span_y0
            current_y1 = span_y1

    # Flush the last line
    if current_line_spans:
        current_line_spans.sort(key=lambda x: x[0][0])
        joined_text = "  ".join([s[1] for s in current_line_spans]).strip()
        if joined_text:
            lines.append(joined_text)

    return "\n".join(lines) + "\n"


def extract_text_from_pdf_sync(file_path: str) -> str:
    """
    Synchronous version of PDF text extraction using PyMuPDF.
    Sorts blocks by vertical (Y) coordinate to keep dates aligned.
    """
    print("Starting text extraction from PDF (sync).")
    document = fitz.open(file_path)
    text = ""
    print(f"PDF has {len(document)} page(s).")
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += _extract_page_text_by_y(page)
    logging.info("Text extraction from PDF completed.")
    print("=================PDF TEXT IS EXTRACTED=====================")
    return text


async def extract_text_from_pdf_stream(file) -> str:
    """
    Extract text from a PDF file object (in-memory bytes stream).
    """
    print("Starting text extraction from PDF stream.")
    file_content = await file.read()
    document = fitz.open(stream=file_content, filetype="pdf")
    text = ""
    print(f"PDF has {len(document)} page(s).")
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += _extract_page_text_by_y(page)
    logging.info("Text extraction from PDF completed.")
    print("=================PDF TEXT IS EXTRACTED=====================")
    return text


def pdf_has_text(file_path: str) -> bool:
    """
    Check whether a PDF contains selectable text (i.e., is not a scanned/image PDF).
    Returns True if text is found on consecutive pages (or the single page).
    """
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)

        # Single-page document
        if len(reader.pages) == 1:
            page = reader.pages[0]
            text = page.extract_text()
            return bool(text and text.strip())

        # Multi-page: check consecutive pages for text
        for page_num in range(len(reader.pages) - 1):
            current_page = reader.pages[page_num]
            next_page = reader.pages[page_num + 1]

            current_text = current_page.extract_text()
            next_text = next_page.extract_text()

            if current_text and current_text.strip() and next_text and next_text.strip():
                return True

        return False


# ---------------------------------------------------------------------------
# Word (DOCX) Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_word(file_path: str) -> str:
    """
    Extract text from a Word document (.docx), preserving numbered list formatting.
    Accepts a file path string.
    """
    doc = Document(file_path)
    text = ""
    counters = {}  # {(numId, ilvl): count}

    for para in doc.paragraphs:
        pPr   = getattr(para._p, "pPr", None)
        numPr = getattr(pPr, "numPr", None) if pPr else None

        if numPr:
            numId = int(numPr.numId.val)
            ilvl  = int(numPr.ilvl.val)
            key   = (numId, ilvl)
            counters[key] = counters.get(key, 0) + 1

            if ilvl == 0:
                label = f"{counters[key]}. "
            else:
                label = f"{chr(ord('a') + counters[key] - 1)}. "

            text += label + para.text + "\n"
        else:
            text += para.text + "\n"
    
    print("=================DOCX TEXT IS EXTRACTED=====================")
    return text


# ---------------------------------------------------------------------------
# Word (DOCX) Document Generation
# ---------------------------------------------------------------------------

def save_text_to_docx(text: str, output_path: str = None) -> str:
    """
    Save plain text to a Word document (.docx).
    Each non-empty line becomes a paragraph.
    Returns the output file path.
    """
    print("Saving extracted text to Word document...")
    
    if output_path is None:
        unique_id = uuid.uuid4()
        output_path = os.path.join('/tmp', f"input_{str(unique_id)}.docx")

    doc = Document()
    for line in text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(output_path)

    print(f"Word document saved at: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Document Correction & Redlining
# ---------------------------------------------------------------------------

def apply_corrections_for_clauses(file_path: str, differences: list, suggestions: list, output_path: str) -> str:
    """
    Apply clause-level corrections to a Word document.
    
    Args:
        file_path: Path to the original .docx file.
        differences: List of dicts with 'Uploaded NDA clause' keys.
        suggestions: List of dicts with 'Correction' and 'Change' keys.
        output_path: Path to save the corrected document.
    
    Returns:
        The output file path.
    """
    # Filter only changes flagged for application
    filtered = [
        (diff, sugg)
        for diff, sugg in zip(differences, suggestions)
        if sugg.get("Change") is True
    ]
    print(f"\nApplying {len(filtered)} corrections out of {len(suggestions)} suggestions\n")

    doc = Document(file_path)
    paras = doc.paragraphs

    print("\n======== [DOC PARAGRAPHS] ========")
    for i, p in enumerate(paras):
        txt = p.text.replace("\n", "\\n")
        print(f"[{i:03d}] '{txt}'")
    print("=======================================\n")

    diff_idx = 0
    total = len(filtered)

    while diff_idx < total:
        uploaded_raw = filtered[diff_idx][0]["Uploaded NDA clause"].strip()
        correction = filtered[diff_idx][1]["Correction"].strip()

        start_hint = ' '.join(uploaded_raw.lower().split()[:6])
        end_hint = ' '.join(uploaded_raw.lower().split()[-6:])

        print(f"Correction #{diff_idx + 1}")
        print(f"Start hint: '{start_hint}'")
        print(f"End hint  : '{end_hint}'")

        # Search for matching paragraphs in the document
        start_idx, end_idx = None, None
        for i, para in enumerate(paras):
            norm = normalize(para.text)
            if start_idx is None and start_hint in norm:
                start_idx = i
                print(f"Found start at paragraph {i}")
            if start_idx is not None and end_hint in norm:
                end_idx = i
                print(f"Found end at paragraph {i}")
                break

        # Sliding window search if end not found in same paragraph
        if start_idx is not None and end_idx is None:
            print("\nEnd not found in the same paragraph — scanning forward by window\n")
            concated_text = normalize(paras[start_idx].text)
            j = start_idx + 1
            while j < len(paras):
                concated_text += " " + normalize(paras[j].text)
                if end_hint in concated_text:
                    end_idx = j
                    print(f"[Search by window] Found end across paragraphs up to {j}")
                    break
                j += 1
            if end_idx is None:
                print("[Search by window] Reached end of document — end hint NOT found.")

        if start_idx is not None and end_idx is not None:
            print(f"Treating paragraphs {start_idx} to {end_idx} as one block.\n")

            # Preserve text before the clause in the same paragraph
            original_start_text = paras[start_idx].text
            start_pos = normalize(original_start_text).find(start_hint)
            if start_pos != -1:
                before_start = original_start_text[:original_start_text.lower().find(start_hint)]
                print(f"Keeping heading: '{before_start.strip()}'")
            else:
                before_start = ""

            parts = [p.text for p in paras[start_idx:end_idx + 1]]
            block_text = " ".join(parts)
            new_text = correction
            
            if flatten(uploaded_raw) in flatten(block_text):
                paras[start_idx].text = block_text.replace(uploaded_raw, new_text, 1)
                print("[REPLACE] literal match: True\n", "="*40)
            else:
                paras[start_idx].text = before_start + new_text
                print("[REPLACE] literal match: False so replace the correction manually\n", "@"*40)

            for j in range(end_idx, start_idx, -1):
                p = paras[j]
                p._element.getparent().remove(p._element)

            # Refresh paragraph list after deletion
            paras = doc.paragraphs

        else:
            print(f"Could not find match for correction #{diff_idx + 1}")

        diff_idx += 1

    doc.save(output_path)
    print(f"\nSaved to {output_path}")
    return output_path


def remove_columns_from_docx(docx_path: str) -> None:
    """
    Remove column layout (<w:cols>) from a Word document's section properties.
    Operates directly on the docx XML.
    """
    temp_dir = mkdtemp()
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    document_path = os.path.join(temp_dir, "word/document.xml")
    tree = etree.parse(document_path)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for sectPr in tree.xpath("//w:sectPr", namespaces=ns):
        cols = sectPr.find("w:cols", namespaces=ns)
        if cols is not None:
            sectPr.remove(cols)

    tree.write(document_path, xml_declaration=True, encoding='utf-8', standalone="yes")

    with zipfile.ZipFile(docx_path, "w") as zipf:
        for foldername, _, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, temp_dir)
                zipf.write(file_path, arcname)

    shutil.rmtree(temp_dir)


def generate_redline(author_tag: str, original_path: str, modified_path: str, output_path: str) -> str:
    """
    Generate a redline (track-changes) comparison between two Word documents.
    
    Args:
        author_tag: Author name to tag the changes with.
        original_path: Path to the original .docx.
        modified_path: Path to the modified .docx.
        output_path: Path to save the redlined .docx.
    
    Returns:
        The output file path.
    """
    engine = XmlPowerToolsEngine()
    original_bytes = Path(original_path).read_bytes()
    modified_bytes = Path(modified_path).read_bytes()

    redline_bytes, stdout_str, stderr_str = engine.run_redline(
        author_tag, original_bytes, modified_bytes
    )

    if stderr_str:
        print("Redline engine output:", stderr_str)

    Path(output_path).write_bytes(redline_bytes)
    remove_columns_from_docx(output_path)
    return output_path
